from tkinter import*
from tkinter import filedialog
from tkinter import ttk
import pandas as pd
import numpy as np
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

    
class Setting:
    def __init__(self, window):
        self.window = window
        window.configure(bg="#EFECE3")  # set any color you want
        #===========Color==============
        self.textColor = StringVar()
        self.BG = StringVar()
        self.darkOrange = StringVar()
        self.buttonBG = StringVar()
        self.buttonText = StringVar()
        self.lightBlur = StringVar()

        self.textColor.set("Black")
        self.BG.set("#EFECE3")
        self.darkOrange.set("#4A70A9")
        self.lightBlur.set("#DCE2F0")
        self.buttonBG.set("#4A70A9")
        self.buttonText.set("Black")

        #===========Language==============
        self.languageOptions = ["English", "Persian", "Hindi"]
        self.languageSelectedOption = StringVar(value=self.languageOptions[0])


class MainActivity:
    def __init__(self, window, setting):
        self.window = window
        self.setting = setting

        frameTitle = Frame(self.window, bg=self.setting.darkOrange.get(), width=400, height=50)
        frameTitle.pack(fill=X)
        labelTitle = Label(frameTitle, text="Easy Report", font=('Arial', 30, 'bold'), bg=self.setting.darkOrange.get())
        labelTitle.pack(side=LEFT, fill=BOTH)

        

        #=========== Section Container ==============
        sectionFrame = Frame(self.window, bg=self.setting.lightBlur.get(), bd=2, relief=RIDGE)
        sectionFrame.pack(fill=X, padx=15, pady=15, ipadx=5, ipady=5)

        #=========== File Direction Chooser ==============
        frameChooser = Frame(sectionFrame, bg=self.setting.lightBlur.get())
        frameChooser.pack(fill=X, padx=20, pady=10)

        labelChooser = Label(frameChooser, text="Student Homework Excel File Location:", bg=self.setting.lightBlur.get(), fg=self.setting.textColor.get(), font=('Arial', 12))
        labelChooser.grid(row=0, column=0, columnspan=2, sticky=W, padx=5, pady=5)

        self.filePath = StringVar()
        entryPath = Entry(frameChooser, textvariable=self.filePath, width=50, bg=self.setting.buttonBG.get())
        entryPath.grid(row=1, column=0, padx=5, pady=5, sticky=W)

        buttonBrowse = Button(frameChooser, text="Browse", bg=self.setting.lightBlur.get(), fg="Black", highlightbackground=self.setting.lightBlur.get(), font=('Arial', 10, 'bold'), command=self.choose_file)
        buttonBrowse.grid(row=1, column=1, padx=5, pady=5, sticky=W)

        #=========== Second File Direction Chooser ==============
        frameChooser2 = Frame(sectionFrame, bg=self.setting.lightBlur.get())
        frameChooser2.pack(fill=X, padx=20, pady=10)

        labelChooser2 = Label(frameChooser2, text="Student Information Excel File Location:", bg=self.setting.lightBlur.get(), fg=self.setting.textColor.get(), font=('Arial', 12))
        labelChooser2.grid(row=0, column=0, columnspan=2, sticky=W, padx=5, pady=5)

        self.studentInfo = StringVar()
        entryPath2 = Entry(frameChooser2, textvariable=self.studentInfo, width=50, bg=self.setting.buttonBG.get())
        entryPath2.grid(row=1, column=0, padx=5, pady=5, sticky=W)

        buttonBrowse2 = Button(frameChooser2, text="Browse", bg=self.setting.lightBlur.get(), fg="Black", highlightbackground=self.setting.lightBlur.get(), font=('Arial', 10, 'bold'), command=self.choose_teacher_file)
        buttonBrowse2.grid(row=1, column=1, padx=5, pady=5, sticky=W)

        #=========== Student Word Report Template Chooser ==============
        frameChooser3 = Frame(sectionFrame, bg=self.setting.lightBlur.get())
        frameChooser3.pack(fill=X, padx=20, pady=10)

        labelChooser3 = Label(frameChooser3, text="Student Word Report Template Location:", bg=self.setting.lightBlur.get(), fg=self.setting.textColor.get(), font=('Arial', 12))
        labelChooser3.grid(row=0, column=0, columnspan=2, sticky=W, padx=5, pady=5)

        self.templatePath = StringVar()
        entryPath3 = Entry(frameChooser3, textvariable=self.templatePath, width=50, bg=self.setting.buttonBG.get())
        entryPath3.grid(row=1, column=0, padx=5, pady=5, sticky=W)

        buttonBrowse3 = Button(frameChooser3, text="Browse", bg=self.setting.lightBlur.get(), fg="Black", highlightbackground=self.setting.lightBlur.get(), font=('Arial', 10, 'bold'), command=self.choose_template_file)
        buttonBrowse3.grid(row=1, column=1, padx=5, pady=5, sticky=W)

        separator = Frame(sectionFrame, bg="black", height=2)
        separator.pack(fill=X, padx=20, pady=10)

        #=========== Save Location Chooser ==============
        frameSave = Frame(sectionFrame, bg=self.setting.lightBlur.get())
        frameSave.pack(fill=X, padx=20, pady=10)

        labelSave = Label(frameSave, text="Save Report Files Location:", bg=self.setting.lightBlur.get(), fg=self.setting.textColor.get(), font=('Arial', 12))
        labelSave.grid(row=0, column=0, columnspan=2, sticky=W, padx=5, pady=5)

        self.savePath = StringVar()
        entrySave = Entry(frameSave, textvariable=self.savePath, width=50, bg=self.setting.buttonBG.get())
        entrySave.grid(row=1, column=0, padx=5, pady=5, sticky=W)

        buttonSave = Button(frameSave, text="Browse", bg=self.setting.lightBlur.get(), fg="Black", highlightbackground=self.setting.lightBlur.get(), font=('Arial', 10, 'bold'), command=self.choose_save_location)
        buttonSave.grid(row=1, column=1, padx=5, pady=5, sticky=W)

        #=========== Action Type Chooser ==============
        #=========== Action Type Section ==============
        actionSection = Frame(self.window, bg=self.setting.BG.get(), bd=2, relief=RIDGE)
        actionSection.pack(fill=X, padx=15, pady=15, ipadx=5, ipady=5)

        frameAction = Frame(actionSection, bg=self.setting.BG.get())
        frameAction.pack(fill=X, padx=20, pady=20)

        labelAction = Label(frameAction, text="Select Action Type:", bg=self.setting.BG.get(), fg=self.setting.textColor.get(), font=('Arial', 12))
        labelAction.grid(row=0, column=0, sticky=W, padx=5, pady=5)

        self.actionOptions = ["All", "Manager Excel Report", "Students Word Report"]
        self.actionSelected = StringVar(value=self.actionOptions[0])

        actionMenu = OptionMenu(frameAction, self.actionSelected, *self.actionOptions)
        actionMenu.config(width=25, font=('Arial', 10), bg=self.setting.BG.get(), fg="black", highlightbackground=self.setting.buttonBG.get())
        actionMenu.grid(row=0, column=1, padx=5, pady=5, sticky=W)

        #=========== Month and Year Chooser ============
        # Month chooser
        months = [
            "January", "February", "March", "April", "May", "June",
            "July", "August", "September", "October", "November", "December"
        ]
        self.selectedMonth = StringVar(value=months[0])
        labelMonth = Label(frameAction, text="Select Report Month:", bg=self.setting.BG.get(), fg=self.setting.textColor.get(), font=('Arial', 12))
        labelMonth.grid(row=1, column=0, sticky=W, padx=5, pady=5)
        monthMenu = OptionMenu(frameAction, self.selectedMonth, *months)
        monthMenu.config(width=15, font=('Arial', 10), bg=self.setting.BG.get(), fg="black", highlightbackground=self.setting.buttonBG.get())
        monthMenu.grid(row=1, column=1, padx=5, pady=5, sticky=W)

        # Year chooser (let's say 2020-2030)
        years = [str(year) for year in range(2024, 2050)]
        self.selectedYear = StringVar(value=years[0])
        labelYear = Label(frameAction, text="Select Report Year:", bg=self.setting.BG.get(), fg=self.setting.textColor.get(), font=('Arial', 12))
        labelYear.grid(row=2, column=0, sticky=W, padx=5, pady=5)
        yearMenu = OptionMenu(frameAction, self.selectedYear, *years)
        yearMenu.config(width=15, font=('Arial', 10), bg=self.setting.BG.get(), fg="black", highlightbackground=self.setting.buttonBG.get())
        yearMenu.grid(row=2, column=1, padx=5, pady=5, sticky=W)

        

        #=========== Progress Bar Section ==============
        progressSection = Frame(self.window, bg=self.setting.BG.get(), bd=2, relief=RIDGE)
        progressSection.pack(fill=X, padx=15, pady=15, ipadx=5, ipady=5)

        labelProgress = Label(progressSection, text="Progress:", bg=self.setting.BG.get(), fg=self.setting.textColor.get(), font=('Arial', 12))
        labelProgress.pack(anchor=W, padx=20, pady=(10, 5))

        self.progress = ttk.Progressbar(progressSection, orient=HORIZONTAL, length=400, mode='determinate')
        self.progress.pack(fill=X, padx=(20,10), pady=10)

        self.progress['value'] = 0

        #=========== Start Button at Right ============
        buttonStart = Button(self.window, text="Start", bg=self.setting.lightBlur.get(), fg="Black", highlightbackground=self.setting.BG.get(), font=('Arial', 12, 'bold'), command=self.start_process)
        buttonStart.pack(side=RIGHT, padx=20, pady=15)

        #=========== Support Button on Left ============
        buttonSupport = Button(self.window, text="Support", bg=self.setting.lightBlur.get(), fg="Black", highlightbackground=self.setting.BG.get(), font=('Arial', 12, 'bold'), command=self.open_support)
        buttonSupport.pack(side=LEFT, padx=20, pady=15)
        
        #=========== Developer Label Between Buttons ============
        labelDeveloper = Label(self.window, text="Developed by Parsa Asil", bg=self.setting.BG.get(), fg=self.setting.textColor.get(), font=('Arial', 12, 'italic'))
        labelDeveloper.pack(side=LEFT, pady=15, padx=(128, 0))

        #=========== Developer Info Section ============
        developerFrame = Frame(self.window, bg=self.setting.BG.get())
        developerFrame.pack(side=RIGHT, padx=(0, 140), pady=15)  # adjust padding to sit left of Start button


    def start_process(self):
        # Example functionality: increase progress bar gradually
        self.progress['value'] = 0
        self.window.update_idletasks()

        if (self.actionSelected.get() == "Manager Excel Report"):
            student_hw_path = self.filePath.get()
            student_info_path = self.studentInfo.get()
            save_path = self.savePath.get()

            report_generator = ManagerReportGenerator(student_hw_path, student_info_path, save_path, self.selectedMonth.get(), self.selectedYear.get())
            report_generator.generate_report()
            print("Manager Excel Report generated successfully at:", save_path)

        elif (self.actionSelected.get() == "Students Word Report"):
            # Construct the path to the saved Manager Excel Report
            save_path = self.savePath.get()
            student_report_folder = os.path.join(
                save_path,
                f"Student_Word_Report_{self.selectedMonth.get()}_{self.selectedYear.get()}"
            )
            os.makedirs(student_report_folder, exist_ok=True)
            manager_report_path = os.path.join(
                self.savePath.get(),
                f"Manager_Report_{self.selectedMonth.get()}_{self.selectedYear.get()}.xlsx"
            )
            generator = StudentReportGenerator(manager_report_path, self.templatePath.get(), student_report_folder)
            generator.generate_reports()

        elif self.actionSelected.get() == "All":
            # First generate Manager Excel Report
            student_hw_path = self.filePath.get()
            student_info_path = self.studentInfo.get()
            save_path = self.savePath.get()

            report_generator = ManagerReportGenerator(student_hw_path, student_info_path, save_path, self.selectedMonth.get(), self.selectedYear.get())
            report_generator.generate_report()
            print("Manager Excel Report generated successfully at:", save_path)

            # Path to Manager Excel Report
            manager_report_path = os.path.join(
                save_path,
                f"Manager_Report_{self.selectedMonth.get()}_{self.selectedYear.get()}.xlsx"
            )

            # Create a new folder for student Word reports
            student_report_folder = os.path.join(
                save_path,
                f"Student_Word_Report_{self.selectedMonth.get()}_{self.selectedYear.get()}"
            )
            os.makedirs(student_report_folder, exist_ok=True)

            # Generate Student Word Reports
            generator = StudentReportGenerator(manager_report_path, self.templatePath.get(), student_report_folder)
            generator.generate_reports()



        for i in range(0, 101, 10):
            self.progress['value'] = i
            self.window.update()
            self.window.after(100)  # 0.1 second delay


    def choose_file(self):
        file_path = filedialog.askopenfilename(
            title="Student Homework Excel File",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        if file_path:
            self.filePath.set(file_path)
        print(file_path)

    def choose_teacher_file(self):
        file_path = filedialog.askopenfilename(
            title="Student Information Excel File",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        if file_path:
            self.studentInfo.set(file_path)
        print(file_path)

    def choose_template_file(self):
        file_path = filedialog.askopenfilename(
            title="Student Word Report Template",
            filetypes=[("Word files", "*.docx")]
        )
        if file_path:
            self.templatePath.set(file_path)
        print(file_path)

    def choose_save_location(self):
        folder_path = filedialog.askdirectory(
            title="Select Folder to Save File"
        )
        if folder_path:
            self.savePath.set(folder_path)
            self.selected_save_folder = folder_path

    def open_support(self):
        clearWindow(self.window)
        SuppportPage(self.window)


class ManagerReportGenerator:
    def __init__(self, studentHW, student_info, save_path=None, month=None, year=None):
        """
        Initialize StudentReportGenerator with paths.
        """
        self.studentHW = studentHW
        self.student_info = student_info
        self.save_path = save_path
        self.month = month
        self.year = year

    def generate_report(self):
        """
        Generate formatted student report from master homework Excel file and student info file.
        Returns:
            pd.DataFrame: Final formatted report DataFrame
        """
        teacher = pd.read_excel(self.studentHW)
        # Filter data by selected month and year based on "Date Received"
        teacher['Date Received'] = pd.to_datetime(teacher['Date Received'], errors='coerce')
        teacher = teacher[
            (teacher['Date Received'].dt.month_name() == self.month) &
            (teacher['Date Received'].dt.year == int(self.year))
        ]

        grade_to_num = {
            "A+": 4.0, "A": 3.7, "A-": 3.3,
            "B+": 3.0, "B": 2.7, "B-": 2.3,
            "C+": 2.0, "C": 1.7, "C-": 1.3,
            "D+": 1.0, "D": 0.7, "D-": 0.3,
            "F": 0
        }

        num_to_grade = {v: k for k, v in grade_to_num.items()}

        teacher["Grade_Num"] = teacher["Grade"].map(grade_to_num)

        aggregated = teacher.groupby(["Student Name", "Topic Name", "Sub Topic"], as_index=False).agg({
            "Date Issued": "min",
            "Date Received": "max",
            "Booklet No.": "count",
            "Incorrect Answers": "sum",
            "Grade_Num": "mean",
            "Remark": lambda x: ", ".join(x)
        })

        def numeric_to_letter_plusminus(n):
            closest = min(num_to_grade.keys(), key=lambda x: abs(x - n))
            return num_to_grade[closest]

        aggregated["Grade"] = aggregated["Grade_Num"].apply(numeric_to_letter_plusminus)
        aggregated = aggregated.drop(columns=["Grade_Num"])

        def format_student_data(df):
            formatted_rows = []
            for student, group in df.groupby("Student Name"):
                row = {
                    "Student Name": student,
                    "Invoice Number": "",
                    "EmailSubject": "",
                    "Email": "",
                    "bcc": "",
                    "Foot Notes": ""
                }

                topic_idx = 1
                for _, topic_group in group.groupby("Topic Name"):
                    topic_name = topic_group["Topic Name"].iloc[0]
                    subtopic_name = topic_group["Sub Topic"].iloc[0]
                    avg_grade = topic_group["Grade"].iloc[0]
                    row[f"Topic{topic_idx}"] = f"{topic_name}: {subtopic_name}"
                    row[f"Grade{topic_idx}"] = avg_grade
                    topic_idx += 1

                formatted_rows.append(row)

            return pd.DataFrame(formatted_rows)

        final_df = format_student_data(aggregated)

        student_info_df = pd.read_excel(self.student_info)
        final_df = final_df.drop(columns=["Invoice Number", "Email"], errors="ignore")
        final_df = pd.merge(final_df, student_info_df[["Student Name", "Invoice Number", "Email"]], on="Student Name", how="left")

        cols = list(final_df.columns)
        ordered_cols = ["Student Name", "Invoice Number", "Email"] + [c for c in cols if c not in ["Student Name", "Invoice Number", "Email"]]
        final_df = final_df[ordered_cols]

        # Ensure exactly one space between trimmed month and year
        final_df["GradeForTheDateof"] = f"{str(self.month).strip()} {str(self.year).strip()}"

        if self.save_path:
            if not self.save_path.lower().endswith(".xlsx"):
                self.save_path = os.path.join(self.save_path, f"Manager_Report_{self.month}_{self.year}.xlsx")
            final_df.to_excel(self.save_path, index=False)

        return final_df
    

class StudentReportGenerator:
    def __init__(self, excel_path, template_path, output_folder):
        self.excel_path = excel_path
        self.template_path = template_path
        self.output_folder = output_folder
        os.makedirs(self.output_folder, exist_ok=True)

    def replace_placeholder_with_bold_and_size(self, paragraph, placeholder, replacement_text, font_size_pt):
        # Clear existing runs
        for run in paragraph.runs:
            run.clear()
        # Remove all runs
        while paragraph.runs:
            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
        # Add new run with formatted text
        run = paragraph.add_run()
        run.bold = True
        run.font.size = Pt(font_size_pt)
        run.text = replacement_text

    def replace_placeholder_in_cell(self, cell, placeholder, replacement_text, font_size_pt):
        for paragraph in cell.paragraphs:
            if placeholder in paragraph.text:
                # Clear existing runs
                for run in paragraph.runs:
                    run.clear()
                while paragraph.runs:
                    paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
                run = paragraph.add_run()
                run.bold = True
                run.font.size = Pt(font_size_pt)
                run.text = replacement_text

    def insert_table_at_paragraph(self, paragraph, topics_and_grades):
        doc = paragraph.part.document
        parent = paragraph._element.getparent()
        p_idx = parent.index(paragraph._element)

        # Insert table after the paragraph
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'  # Apply default built-in Word table style

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Topic'
        hdr_cells[1].text = 'Average Grade'

        # Apply shading and formatting to header row
        shading_elm_1 = OxmlElement('w:shd')
        shading_elm_1.set(qn('w:fill'), 'D9D9D9')  # Light gray shading
        hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

        shading_elm_2 = OxmlElement('w:shd')
        shading_elm_2.set(qn('w:fill'), 'D9D9D9')
        hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)

        for idx, cell in enumerate(hdr_cells):
            for paragraph_cell in cell.paragraphs:
                for run in paragraph_cell.runs:
                    run.bold = True
                    run.font.size = Pt(16)
                # Center-align header text
                paragraph_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for topic, grade in topics_and_grades:
            row_cells = table.add_row().cells
            row_cells[0].text = topic
            row_cells[1].text = grade
            # Center-align topic and grade text
            for paragraph_cell in row_cells[0].paragraphs:
                for run in paragraph_cell.runs:
                    run.bold = False
                    run.font.size = Pt(14)
                paragraph_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph_cell in row_cells[1].paragraphs:
                for run in paragraph_cell.runs:
                    run.bold = False
                    run.font.size = Pt(14)
                paragraph_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Move the table to the position after the paragraph
        tbl_element = table._element
        parent.insert(p_idx + 1, tbl_element)

        # Clear the paragraph text instead of removing the paragraph element
        paragraph.clear()

    def generate_reports(self):
        # Load Excel
        df = pd.read_excel(self.excel_path)

        # Load Word template
        template = Document(self.template_path)

        # Generate report for each student
        for index, row in df.iterrows():
            student_name = row['Student Name']  # Make sure this matches your Excel column name
            foot_notes = str(row['Foot Notes']) if 'Foot Notes' in row and pd.notna(row['Foot Notes']) else ''
            doc = Document(self.template_path)  # Load a fresh copy for each student

            # Replace placeholder in paragraphs
            for paragraph in doc.paragraphs:
                if "«Student_Name»" in paragraph.text:
                    self.replace_placeholder_with_bold_and_size(paragraph, "«Student_Name»", student_name, 16)
                if "«Foot_Notes»" in paragraph.text:
                    # Replace placeholder with foot notes text preserving formatting
                    for run in paragraph.runs:
                        if "«Foot_Notes»" in run.text:
                            run.text = run.text.replace("«Foot_Notes»", foot_notes)

            # Replace placeholder in tables
            for table in doc.tables:
                for row_cells in table.rows:
                    for cell in row_cells.cells:
                        if "«Student_Name»" in cell.text:
                            self.replace_placeholder_in_cell(cell, "«Student_Name»", student_name, 16)
                        for paragraph in cell.paragraphs:
                            if "«Foot_Notes»" in paragraph.text:
                                for run in paragraph.runs:
                                    if "«Foot_Notes»" in run.text:
                                        run.text = run.text.replace("«Foot_Notes»", foot_notes)

            # Find all topic and grade columns
            topic_cols = [col for col in df.columns if col.startswith('Topic')]
            grade_cols = [col for col in df.columns if col.startswith('Grade')]
            # Sort columns to ensure matching order
            topic_cols.sort()
            grade_cols.sort()

            # Create list of non-empty topics and grades
            topics_and_grades = []
            for t_col, g_col in zip(topic_cols, grade_cols):
                topic = row[t_col]
                if pd.notna(topic) and str(topic).strip() != '':
                    grade = row[g_col]
                    topics_and_grades.append((str(topic), str(grade) if pd.notna(grade) else ''))

            # Replace «Topic_Grades» placeholder with table in paragraphs
            for paragraph in list(doc.paragraphs):
                if "«Topic_Grades»" in paragraph.text:
                    self.insert_table_at_paragraph(paragraph, topics_and_grades)

            # Replace «Topic_Grades» placeholder with table in table cells
            for table in doc.tables:
                for row_cells in table.rows:
                    for cell in row_cells.cells:
                        for paragraph in list(cell.paragraphs):
                            if "«Topic_Grades»" in paragraph.text:
                                self.insert_table_at_paragraph(paragraph, topics_and_grades)

            # Save the new document
            output_path = os.path.join(self.output_folder, f"{student_name}_Report.docx")
            doc.save(output_path)

        print("Reports generated successfully!")


class SuppportPage:
    def __init__(self, window):
        self.window = window
        self.setting = setting

        # Title frame
        frameTitle = Frame(self.window, bg=self.setting.darkOrange.get(), width=400, height=50)
        frameTitle.pack(fill=X)
        labelTitle = Label(frameTitle, text="Support", font=('Arial', 30, 'bold'), bg=self.setting.darkOrange.get(), fg="white")
        labelTitle.pack(side=LEFT, fill=BOTH, padx=10, pady=5)

        # Support Information Frame
        supportFrame = Frame(self.window, bg=self.setting.lightBlur.get(), bd=2, relief=RIDGE)
        supportFrame.pack(fill=BOTH, expand=True, padx=20, pady=20)

        labelInfoTitle = Label(supportFrame, text="Contact Information", font=('Arial', 16, 'bold'), bg=self.setting.lightBlur.get(),fg=self.setting.textColor.get())
        labelInfoTitle.pack(pady=(10, 15))

        # Email
        labelEmail = Label(supportFrame, text="📧 Email: parsa.asil@outlook.com", font=('Arial', 14), bg=self.setting.lightBlur.get(), fg=self.setting.textColor.get())
        labelEmail.pack(anchor=W, padx=20, pady=5)

        # Phone
        labelPhone = Label(supportFrame, text="📞 Phone: +1 832 840 3008", font=('Arial', 14), bg=self.setting.lightBlur.get(),fg=self.setting.textColor.get())
        labelPhone.pack(anchor=W, padx=20, pady=5)

        # Working Hours
        labelHours = Label(supportFrame, text="⏰ Working Hours: Mon-Fri, 9:00 AM - 6:00 PM", font=('Arial', 14), bg=self.setting.lightBlur.get(),fg=self.setting.textColor.get())
        labelHours.pack(anchor=W, padx=20, pady=5)

        # Additional instructions or FAQs
        labelExtra = Label(supportFrame, text="For any issues regarding Excel or Word report generation,\nplease contact support with your report file.", font=('Arial', 12), bg=self.setting.lightBlur.get(),fg=self.setting.textColor.get(), justify=LEFT)
        labelExtra.pack(anchor=W, padx=20, pady=(15, 10))

        # Back Button
        buttonBack = Button(self.window, text="Back", bg=self.setting.lightBlur.get(), fg="Black", highlightbackground=self.setting.BG.get(), font=('Arial', 12, 'bold'), command=self.open_back)
        buttonBack.pack(side=LEFT, padx=20, pady=15)

    def open_back(self):
        clearWindow(self.window)
        MainActivity(self.window, self.setting)

    


def clearWindow(window):
    for widget in window.winfo_children():
        widget.destroy()




window = Tk()
window.title('Easy Report')
window.geometry("630x850")
window.resizable(False, False)

setting = Setting(window)
MainActivity(window,setting)

window.mainloop()